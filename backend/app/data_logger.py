"""
Data Logger for Industrial Digital Twin
Logs telemetry every 10 seconds and generates CSV batch reports
"""

import os
import csv
import json
from datetime import datetime
from collections import deque
from typing import Optional, Dict, Any


class DataLogger:
    """
    Industrial Data Logger for DCP Batch Reactor
    
    Features:
    - Logs telemetry every 10 seconds
    - Stores: timestamp, step number, temperature, flow rate, control actions
    - Generates CSV batch reports
    - In-memory buffer for real-time access
    """
    
    def __init__(self, log_dir: str = "logs", buffer_size: int = 3600):
        self.log_dir = log_dir
        self.buffer_size = buffer_size  # Store last hour at 10s intervals
        
        # In-memory buffer for recent data
        self.telemetry_buffer = deque(maxlen=buffer_size)
        self.event_buffer = deque(maxlen=500)
        
        # Batch tracking
        self.current_batch_id: Optional[str] = None
        self.batch_data: list = []
        self.batch_start_time: Optional[datetime] = None
        
        # Logging interval (10 seconds)
        self.log_interval_seconds = 10
        self.last_log_time = 0
        
        # Control action tracking
        self.control_actions_log = deque(maxlen=100)
        
        # Ensure log directory exists
        os.makedirs(log_dir, exist_ok=True)
    
    def start_batch(self, batch_id: Optional[str] = None) -> str:
        """Start a new batch logging session"""
        self.current_batch_id = batch_id or datetime.now().strftime("BATCH_%Y%m%d_%H%M%S")
        self.batch_data = []
        self.batch_start_time = datetime.now()
        
        # Log batch start event
        self.log_event("INFO", f"Batch started: {self.current_batch_id}")
        
        return self.current_batch_id
    
    def end_batch(self) -> str:
        """End the current batch and generate report"""
        if not self.current_batch_id:
            return ""
        
        # Generate CSV report
        report_path = self.generate_csv_report()
        
        # Log batch end event
        self.log_event("INFO", f"Batch completed: {self.current_batch_id}")
        
        # Reset batch tracking
        batch_id = self.current_batch_id
        self.current_batch_id = None
        self.batch_start_time = None
        
        return report_path
    
    def log_telemetry(self, telemetry: Dict[str, Any], force: bool = False) -> bool:
        """
        Log telemetry data (every 10 seconds or when forced)
        
        Returns True if data was logged, False if skipped
        """
        import time
        current_time = time.time()
        
        # Check if enough time has passed since last log
        if not force and (current_time - self.last_log_time) < self.log_interval_seconds:
            return False
        
        self.last_log_time = current_time
        
        # Create log entry
        entry = {
            "timestamp": datetime.now().isoformat(),
            "unix_time": current_time,
            "step_number": telemetry.get("recipe_step_index", 0),
            "step_name": telemetry.get("recipe_step_name", "N/A"),
            "temperature_c": telemetry.get("temp", 0),
            "pressure_bar": telemetry.get("pressure", 0),
            "cl2_flow_pct": telemetry.get("inputs", {}).get("cl2", 0),
            "cooling_pct": telemetry.get("inputs", {}).get("cooling", 0),
            "agitator_rpm": telemetry.get("inputs", {}).get("rpm", 0),
            "discharge_pct": telemetry.get("inputs", {}).get("discharge", 0),
            "purity_pct": telemetry.get("purity", 0),
            "moles_dcp": telemetry.get("moles_dcp", 0),
            "moles_phenol": telemetry.get("moles_phenol", 0),
            "financial_value": telemetry.get("financial_value", 0),
            "safety_scram": telemetry.get("safety_scram", False),
            "recipe_status": telemetry.get("recipe_status", ""),
            "interlock_active": telemetry.get("interlock_active", False),
            "interlock_msg": telemetry.get("interlock_msg", "")
        }
        
        # Add to buffer
        self.telemetry_buffer.append(entry)
        
        # Add to batch data if batch is active
        if self.current_batch_id:
            self.batch_data.append(entry)
        
        return True
    
    def log_control_action(self, action_type: str, value: Any, source: str = "operator"):
        """Log a control action"""
        entry = {
            "timestamp": datetime.now().isoformat(),
            "action_type": action_type,
            "value": value,
            "source": source
        }
        
        self.control_actions_log.append(entry)
        
        # Also add to batch data if batch is active
        if self.current_batch_id:
            # Add as a special entry type
            action_entry = {
                "timestamp": entry["timestamp"],
                "entry_type": "control_action",
                "action": action_type,
                "value": str(value),
                "source": source
            }
            self.batch_data.append(action_entry)
    
    def log_event(self, level: str, message: str, meta: Optional[Dict] = None):
        """Log an event"""
        entry = {
            "timestamp": datetime.now().isoformat(),
            "level": level,
            "message": message,
            "meta": meta or {}
        }
        
        self.event_buffer.append(entry)
    
    def generate_csv_report(self, batch_id: Optional[str] = None) -> str:
        """
        Generate a CSV batch report
        
        Format:
        - Timestamp
        - Step Number
        - Temperature
        - Flow Rate (Cl2)
        - Control Actions
        """
        target_batch_id = batch_id or self.current_batch_id
        if not target_batch_id:
            target_batch_id = datetime.now().strftime("REPORT_%Y%m%d_%H%M%S")
        
        filename = f"{target_batch_id}.csv"
        filepath = os.path.join(self.log_dir, filename)
        
        # Get data to export
        data_to_export = self.batch_data if self.batch_data else list(self.telemetry_buffer)
        
        if not data_to_export:
            return ""
        
        # Write CSV
        headers = [
            "Timestamp",
            "Step_Number",
            "Step_Name",
            "Temperature_C",
            "Pressure_Bar",
            "Cl2_Flow_Pct",
            "Cooling_Pct",
            "Agitator_RPM",
            "Purity_Pct",
            "Moles_DCP",
            "Financial_Value",
            "Recipe_Status",
            "Safety_Scram",
            "Interlock_Active"
        ]
        
        try:
            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                
                # Write header
                writer.writerow(headers)
                
                # Write data rows (only telemetry entries, not control actions)
                for entry in data_to_export:
                    if entry.get("entry_type") == "control_action":
                        continue  # Skip control action entries for main CSV
                    
                    row = [
                        entry.get("timestamp", ""),
                        entry.get("step_number", 0),
                        entry.get("step_name", ""),
                        entry.get("temperature_c", 0),
                        entry.get("pressure_bar", 0),
                        entry.get("cl2_flow_pct", 0),
                        entry.get("cooling_pct", 0),
                        entry.get("agitator_rpm", 0),
                        entry.get("purity_pct", 0),
                        entry.get("moles_dcp", 0),
                        entry.get("financial_value", 0),
                        entry.get("recipe_status", ""),
                        entry.get("safety_scram", False),
                        entry.get("interlock_active", False)
                    ]
                    writer.writerow(row)
            
            print(f"[LOGGER] CSV report generated: {filepath}")
            return filepath
        
        except Exception as e:
            print(f"[LOGGER] Error generating CSV: {e}")
            return ""
    
    def generate_batch_summary(self) -> Dict[str, Any]:
        """Generate a summary of the current/last batch"""
        if not self.batch_data:
            return {"error": "No batch data available"}
        
        # Calculate statistics
        temps = [e.get("temperature_c", 0) for e in self.batch_data if "temperature_c" in e]
        pressures = [e.get("pressure_bar", 0) for e in self.batch_data if "pressure_bar" in e]
        
        summary = {
            "batch_id": self.current_batch_id or "Unknown",
            "start_time": self.batch_data[0].get("timestamp") if self.batch_data else None,
            "end_time": self.batch_data[-1].get("timestamp") if self.batch_data else None,
            "total_records": len(self.batch_data),
            "duration_seconds": (len(self.batch_data) * self.log_interval_seconds),
            "temperature": {
                "min": min(temps) if temps else 0,
                "max": max(temps) if temps else 0,
                "avg": sum(temps) / len(temps) if temps else 0
            },
            "pressure": {
                "min": min(pressures) if pressures else 0,
                "max": max(pressures) if pressures else 0,
                "avg": sum(pressures) / len(pressures) if pressures else 0
            },
            "final_purity": self.batch_data[-1].get("purity_pct", 0) if self.batch_data else 0,
            "final_dcp_yield": self.batch_data[-1].get("moles_dcp", 0) if self.batch_data else 0,
            "final_value": self.batch_data[-1].get("financial_value", 0) if self.batch_data else 0,
            "safety_events": sum(1 for e in self.batch_data if e.get("safety_scram")),
            "interlock_events": sum(1 for e in self.batch_data if e.get("interlock_active"))
        }
        
        return summary
    
    def get_recent_telemetry(self, count: int = 100) -> list:
        """Get recent telemetry entries"""
        return list(self.telemetry_buffer)[-count:]
    
    def get_recent_events(self, count: int = 50) -> list:
        """Get recent events"""
        return list(self.event_buffer)[-count:]
    
    def get_control_actions(self, count: int = 50) -> list:
        """Get recent control actions"""
        return list(self.control_actions_log)[-count:]

    def generate_pdf_report(self, batch_id: Optional[str] = None) -> str:
        """Generate a PDF batch report using fpdf2."""
        try:
            from fpdf import FPDF
        except ImportError:
            print("[LOGGER] fpdf2 not installed")
            return ""

        target_batch_id = batch_id or self.current_batch_id
        if not target_batch_id:
            target_batch_id = datetime.now().strftime("REPORT_%Y%m%d_%H%M%S")

        filepath = os.path.join(self.log_dir, f"{target_batch_id}.pdf")
        data = self.batch_data if self.batch_data else list(self.telemetry_buffer)
        summary = self.generate_batch_summary()

        try:
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            # Title
            pdf.set_font("Helvetica", "B", 18)
            pdf.cell(0, 12, "DCP-TWIN PRO  -  Batch Report", new_x="LMARGIN", new_y="NEXT", align="C")
            pdf.set_font("Helvetica", "", 10)
            pdf.cell(0, 6, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", new_x="LMARGIN", new_y="NEXT", align="C")
            pdf.ln(6)

            # Batch Summary
            pdf.set_font("Helvetica", "B", 13)
            pdf.cell(0, 8, "Batch Summary", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

            pdf.set_font("Helvetica", "B", 9)
            pdf.set_fill_color(40, 40, 60)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(60, 7, "Parameter", border=1, fill=True)
            pdf.cell(60, 7, "Value", border=1, fill=True, new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 9)

            summary_rows = [
                ("Batch ID", str(summary.get("batch_id", "N/A"))),
                ("Start Time", str(summary.get("start_time", "N/A"))),
                ("End Time", str(summary.get("end_time", "N/A"))),
                ("Total Records", str(summary.get("total_records", 0))),
                ("Duration (s)", str(summary.get("duration_seconds", 0))),
                ("Avg Temperature (C)", f"{summary.get('temperature', {}).get('avg', 0):.1f}"),
                ("Min Temperature (C)", f"{summary.get('temperature', {}).get('min', 0):.1f}"),
                ("Max Temperature (C)", f"{summary.get('temperature', {}).get('max', 0):.1f}"),
                ("Avg Pressure (bar)", f"{summary.get('pressure', {}).get('avg', 0):.2f}"),
                ("Final Purity (%)", f"{summary.get('final_purity', 0):.1f}"),
                ("Final DCP Yield (mol)", f"{summary.get('final_dcp_yield', 0):.4f}"),
                ("Final Value ($)", f"{summary.get('final_value', 0):.2f}"),
                ("Safety Events", str(summary.get("safety_events", 0))),
                ("Interlock Events", str(summary.get("interlock_events", 0))),
            ]
            for label, val in summary_rows:
                pdf.cell(60, 6, label, border=1)
                pdf.cell(60, 6, val, border=1, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(6)

            # Telemetry Data Table
            if data:
                pdf.set_font("Helvetica", "B", 13)
                pdf.cell(0, 8, "Telemetry Data", new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)

                headers = ["Timestamp", "Step", "Temp (C)", "Press", "Cl2 (%)", "Purity", "Status"]
                col_w = [38, 14, 20, 20, 18, 20, 30]
                pdf.set_font("Helvetica", "B", 7)
                pdf.set_fill_color(40, 40, 60)
                pdf.set_text_color(255, 255, 255)
                for i, h in enumerate(headers):
                    pdf.cell(col_w[i], 6, h, border=1, fill=True)
                pdf.ln()
                pdf.set_text_color(0, 0, 0)
                pdf.set_font("Helvetica", "", 7)

                for entry in data[:500]:
                    if entry.get("entry_type") == "control_action":
                        continue
                    ts = str(entry.get("timestamp", ""))[-8:]
                    row = [
                        ts,
                        str(entry.get("step_number", "")),
                        f"{entry.get('temperature_c', 0):.1f}",
                        f"{entry.get('pressure_bar', 0):.2f}",
                        f"{entry.get('cl2_flow_pct', 0):.0f}",
                        f"{entry.get('purity_pct', 0):.1f}",
                        str(entry.get("recipe_status", ""))[:12],
                    ]
                    for i, val in enumerate(row):
                        pdf.cell(col_w[i], 5, val, border=1)
                    pdf.ln()

            pdf.output(filepath)
            print(f"[LOGGER] PDF report generated: {filepath}")
            return filepath
        except Exception as e:
            print(f"[LOGGER] Error generating PDF: {e}")
            return ""

    def generate_docx_report(self, batch_id: Optional[str] = None) -> str:
        """Generate a Word (.docx) batch report using python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("[LOGGER] python-docx not installed")
            return ""

        target_batch_id = batch_id or self.current_batch_id
        if not target_batch_id:
            target_batch_id = datetime.now().strftime("REPORT_%Y%m%d_%H%M%S")

        filepath = os.path.join(self.log_dir, f"{target_batch_id}.docx")
        data = self.batch_data if self.batch_data else list(self.telemetry_buffer)
        summary = self.generate_batch_summary()

        try:
            doc = Document()
            title = doc.add_heading("DCP-TWIN PRO \u2014 Batch Report", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Batch Summary
            doc.add_heading("Batch Summary", level=1)
            summary_rows = [
                ("Batch ID", str(summary.get("batch_id", "N/A"))),
                ("Start Time", str(summary.get("start_time", "N/A"))),
                ("End Time", str(summary.get("end_time", "N/A"))),
                ("Total Records", str(summary.get("total_records", 0))),
                ("Duration (s)", str(summary.get("duration_seconds", 0))),
                ("Avg Temperature (\u00b0C)", f"{summary.get('temperature', {}).get('avg', 0):.1f}"),
                ("Min Temperature (\u00b0C)", f"{summary.get('temperature', {}).get('min', 0):.1f}"),
                ("Max Temperature (\u00b0C)", f"{summary.get('temperature', {}).get('max', 0):.1f}"),
                ("Avg Pressure (bar)", f"{summary.get('pressure', {}).get('avg', 0):.2f}"),
                ("Final Purity (%)", f"{summary.get('final_purity', 0):.1f}"),
                ("Final DCP Yield (mol)", f"{summary.get('final_dcp_yield', 0):.4f}"),
                ("Final Value ($)", f"{summary.get('final_value', 0):.2f}"),
                ("Safety Events", str(summary.get("safety_events", 0))),
                ("Interlock Events", str(summary.get("interlock_events", 0))),
            ]
            table = doc.add_table(rows=1, cols=2, style="Light Shading Accent 1")
            table.rows[0].cells[0].text = "Parameter"
            table.rows[0].cells[1].text = "Value"
            for label, val in summary_rows:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = val

            # Telemetry Data
            if data:
                doc.add_heading("Telemetry Data", level=1)
                hdrs = ["Timestamp", "Step", "Temp (\u00b0C)", "Pressure (bar)", "Cl2 (%)", "Purity (%)", "Status"]
                tbl = doc.add_table(rows=1, cols=len(hdrs), style="Light Grid Accent 1")
                for i, h in enumerate(hdrs):
                    tbl.rows[0].cells[i].text = h

                for entry in data[:500]:
                    if entry.get("entry_type") == "control_action":
                        continue
                    row = tbl.add_row()
                    row.cells[0].text = str(entry.get("timestamp", ""))
                    row.cells[1].text = str(entry.get("step_number", ""))
                    row.cells[2].text = f"{entry.get('temperature_c', 0):.1f}"
                    row.cells[3].text = f"{entry.get('pressure_bar', 0):.2f}"
                    row.cells[4].text = f"{entry.get('cl2_flow_pct', 0):.0f}"
                    row.cells[5].text = f"{entry.get('purity_pct', 0):.1f}"
                    row.cells[6].text = str(entry.get("recipe_status", ""))

                for row in tbl.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)

            doc.save(filepath)
            print(f"[LOGGER] DOCX report generated: {filepath}")
            return filepath
        except Exception as e:
            print(f"[LOGGER] Error generating DOCX: {e}")
            return ""


# Global logger instance
data_logger = DataLogger()
